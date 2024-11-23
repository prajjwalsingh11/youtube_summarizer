import os
import nltk
from nltk.data import find
import streamlit as st
from youtube_transcript_api import YouTubeTranscriptApi
from langchain.text_splitter import RecursiveCharacterTextSplitter
from dotenv import load_dotenv
from fpdf import FPDF
import yt_dlp
from docx import Document
from io import BytesIO
import re
import math
from nltk.sentiment.vader import SentimentIntensityAnalyzer
from pytube import YouTube
from pytube.exceptions import PytubeError
from moviepy.video.io.ffmpeg_tools import ffmpeg_extract_subclip
import cv2
import subprocess
from urllib.error import HTTPError
from openai import OpenAI

# Check if the vader_lexicon is already downloaded 
try: 
    find('sentiment/vader_lexicon.zip') 
except LookupError: 
    nltk.download('vader_lexicon')

def load_environment():
    """Load environment variables"""
    env_path = r"C:\coding\youtube summarizer\youtube_summarizer\project.env"
    if os.path.exists(env_path):
        load_dotenv(env_path)
    
    api_key = os.getenv('GROQ_API_KEY')
    if not api_key:
        raise ValueError("GROQ_API_KEY not found in environment variables")
    
    return api_key

# Initialize Groq client
try:
    api_key = load_environment()
    groq_client = OpenAI(
        api_key=api_key,
        base_url="https://api.groq.com/openai/v1"
    )
except Exception as e:
    st.error(f"Error initializing API client: {str(e)}")
    st.stop()

def extract_video_id(youtube_url):
    """Extract video ID from different YouTube URL formats"""
    patterns = [
        r'(?:v=|\/)([0-9A-Za-z_-]{11}).*',  # Standard and shared URLs
        r'(?:embed\/)([0-9A-Za-z_-]{11})',   # Embed URLs
        r'(?:youtu\.be\/)([0-9A-Za-z_-]{11})',  # Shortened URLs
        r'(?:shorts\/)([0-9A-Za-z_-]{11})',   # YouTube Shorts
        r'^([0-9A-Za-z_-]{11})$'  # Just the video ID
    ]
    
    youtube_url = youtube_url.strip()
    
    for pattern in patterns:
        match = re.search(pattern, youtube_url)
        if match:
            return match.group(1)
    
    raise ValueError("Could not extract video ID from URL")

def get_transcript(youtube_url):
    """Get transcript using YouTube Transcript API with cookies"""
    try:
        video_id = extract_video_id(youtube_url)
        
        # Get cookies file path
        cookies_file = os.getenv('COOKIE_PATH', os.path.join(os.path.dirname(__file__), 'cookies.txt'))
        
        if not os.path.exists(cookies_file):
            st.error("Cookie file not found. Please follow the setup instructions in the README.")
            return None, None
            
        try:
            # Read cookies from file
            with open(cookies_file, 'r') as f:
                cookies_content = f.read()
                if not cookies_content.strip():
                    st.error("Cookie file is empty. Please re-export your YouTube cookies.")
                    return None, None
            
            # Get transcript with cookies
            transcript_list = YouTubeTranscriptApi.list_transcripts(video_id, cookies=cookies_file)
            
            try:
                transcript = transcript_list.find_manually_created_transcript()
            except:
                try:
                    transcript = next(iter(transcript_list))
                except Exception as e:
                    st.error("Your YouTube cookies might have expired. Please re-export your cookies and try again.")
                    return None, None
            
            full_transcript = " ".join([part['text'] for part in transcript.fetch()])
            language_code = transcript.language_code
            
            return full_transcript, language_code
                
        except Exception as e:
            st.error("Authentication failed. Please update your cookies.txt file with fresh YouTube cookies.")
            st.info("Tip: Sign in to YouTube again and re-export your cookies using the browser extension.")
            return None, None
            
    except Exception as e:
        st.error("Invalid YouTube URL. Please check the link and try again.")
        return None, None


def get_available_languages():
    """Return a dictionary of available languages"""
    return {
        'English': 'en',
        'Hindi' : 'hi',
        'Deutsch': 'de',
        'Italiano': 'it',
        'Español': 'es',
        'Français': 'fr',
        'Nederlands': 'nl',
        'Polski': 'pl',
        '日本語': 'ja',
        '中文': 'zh',
        'Русский': 'ru',
        '한국어': 'ko',  # Korean
        'Português': 'pt',  # Portuguese
        'العربية': 'ar',  # Arabic
        'Türkçe': 'tr',  # Turkish
        'বাংলা': 'bn',  # Bengali
        'मराठी': 'mr',  # Marathi
        'தமிழ்': 'ta',  # Tamil
        'తెలుగు': 'te',  # Telugu
        'ಕನ್ನಡ': 'kn',  # Kannada
        'മലയാളം': 'ml',  # Malayalam
        'भोजपुरी': 'bh' # Bhojpuri
    }

# Function to create summary prompt
def create_summary_prompt(text, target_language, mode='video'):
    """Create an optimized prompt for summarization in the target language and mode"""
    language_prompts = {
        'en': {
            'title': 'TITLE',
            'overview': 'OVERVIEW',
            'key_points': 'KEY POINTS',
            'takeaways': 'MAIN TAKEAWAYS',
            'context': 'CONTEXT & IMPLICATIONS'
        },
        'hi': {
            'title': 'शीर्षक',
            'overview': 'अवलोकन',
            'key_points': 'मुख्य बिंदु',
            'takeaways': 'मुख्य निष्कर्ष',
            'context': 'प्रसंग और प्रभाव'
        },
        'de': {
            'title': 'TITEL',
            'overview': 'ÜBERBLICK',
            'key_points': 'KERNPUNKTE',
            'takeaways': 'HAUPTERKENNTNISSE',
            'context': 'KONTEXT & AUSWIRKUNGEN'
        },
        'it': { 
            'title': 'TITOLO',
            'overview': 'PANORAMICA',
            'key_points': 'PUNTI CHIAVE',
            'takeaways': 'CONCLUSIONI PRINCIPALI',
            'context': 'CONTESTO E IMPLICAZIONI'
        },
        'es': {
            'title': 'TÍTULO',
            'overview': 'VISIÓN GENERAL',
            'key_points': 'PUNTOS CLAVE',
            'takeaways': 'CONCLUSIONES PRINCIPALES',
            'context': 'CONTEXTO E IMPLICACIONES'
        },
        'fr': {
            'title': 'TITRE',
            'overview': 'APERÇU',
            'key_points': 'POINTS CLÉS',
            'takeaways': 'CONCLUSIONS PRINCIPALES',
            'context': 'CONTEXT & IMPLICATIONS'
        },
        'nl': {
            'title': 'TITEL',
            'overview': 'OVERZICHT',
            'key_points': 'BELANGRIJKE PUNTEN',
            'takeaways': 'HOOFDRESULTATEN',
            'context': 'CONTEXT & IMPLIKATIES'
        },
        'pl': {
            'title': 'TYTUŁ',
            'overview': 'PRZEGLĄD',
            'key_points': 'KLUCZOWE PUNKTY',
            'takeaways': 'GŁÓВНЫЕ ВЫВОДЫ',
            'context': 'КОНТЕКСТ И ИМПЛИКАЦИИ'
        },
        'ja': {
            'title': 'タイトル',
            'overview': '概要',
            'key_points': '主なポイント',
            'takeaways': '主な結論',
            'context': '文脈と影響'
        },
        'zh': {
            'title': '标题',
            'overview': '概述',
            'key_points': '关键点',
            'takeaways': '主要结论',
            'context': '背景与意义'
        },
        'ru': {
            'title': 'ЗАГОЛОВОК',
            'overview': 'ОБЗОР',
            'key_points': 'КЛЮЧЕВЫЕ ПУНКТЫ',
            'takeaways': 'ОСНОВНЫЕ ВЫВОДЫ',
            'context': 'КОНТЕКСТ И ЗНАЧЕНИЕ'
        },
        'ko': {
            'title': '제목',
            'overview': '개요',
            'key_points': '핵심 포인트',
            'takeaways': '주요 결론',
            'context': '맥락 및 의미'
        },
        'pt': {
            'title': 'TÍTULO',
            'overview': 'VISÃO GERAL',
            'key_points': 'PONTOS PRINCIPAIS',
            'takeaways': 'CONCLUSÕES PRINCIPAIS',
            'context': 'CONTEXTO E IMPLICAÇÕES'
        },
        'ar': {
            'title': 'العنوان',
            'overview': 'نظرة عامة',
            'key_points': 'النقاط الرئيسية',
            'takeaways': 'الاستنتاجات الرئيسية',
            'context': 'السياق والآثار'
        },
        'tr': {
            'title': 'BAŞLIK',
            'overview': 'GENEL BAKIŞ',
            'key_points': 'TEMEL NOKTALAR',
            'takeaways': 'ANA SONUÇLAR',
            'context': 'BAĞLAM VE ETKİLER'
        },
        'bn': {
            'title': 'শিরোনাম',
            'overview': 'ওভারভিউ',
            'key_points': 'মূল বিষয়',
            'takeaways': 'মুখ্য বিষয়গুলি',
            'context': 'প্রসঙ্গ ও প্রভাব'
        },
        'mr': {
            'title': 'शीर्षक',
            'overview': 'आढावा',
            'key_points': 'महत्त्वाचे मुद्दे',
            'takeaways': 'मुख्य निष्कर्ष',
            'context': 'संदर्भ आणि परिणाम'
        },
        'ta': {
            'title': 'தலைப்பு',
            'overview': 'அவலோகனம்',
            'key_points': 'முக்கிய அம்சங்கள்',
            'takeaways': 'முக்கிய முடிவுகள்',
            'context': 'சூழல் மற்றும் விளைவுகள்'
        },
        'te': {
            'title': 'శీర్షిక',
            'overview': 'సమగ్రం',
            'key_points': 'ముఖ్య అంశాలు',
            'takeaways': 'ప్రధాన పాయింట్లు',
            'context': 'సందర్భం మరియు ప్రతిఫలాలు'
        },
        'kn': {
            'title': 'ಶೀರ್ಷಿಕೆ',
            'overview': 'ಆವಲೋಕನ',
            'key_points': 'ಮುಖ್ಯ ಅಂಶಗಳು',
            'takeaways': 'ಪ್ರಮುಖ ತತ್ವಗಳು',
            'context': 'ಸಂದರ್ಭ ಮತ್ತು ಪರಿಣಾಮಗಳು'
        },
        'ml': {
            'title': 'തലക്കെട്ട്',
            'overview': 'അവലോകനം',
            'key_points': 'പ്രധാന പോയിന്റുകൾ',
            'takeaways': 'മുഖ്യ സമാഹാരങ്ങൾ',
            'context': 'സന്ദർഭവും പ്രതിഫലങ്ങളും'
        },
        'bh': {
            'title': 'शीर्षक',
            'overview': 'ओवरव्यू',
            'key_points': 'मुख्य बिंदु',
            'takeaways': 'मुख्य निष्कर्ष',
            'context': 'प्रसंग और प्रभाव'
        }
    }

    prompts = language_prompts.get(target_language, language_prompts['en'])

    if mode == 'podcast':
        system_prompt = f"""You are an expert content analyst and summarizer. Create a comprehensive 
        podcast-style summary in {target_language}. Ensure all content is fully translated and culturally adapted 
        to the target language."""

        user_prompt = f"""Please provide a detailed podcast-style summary of the following content in {target_language}. 
        Structure your response as follows, making it an interactive dialogue between two individuals:

        🎙️ {prompts['title']}: Create an engaging title

        🗣️ Interactive Dialogue: 
        - Begin with a question from one person and a response from another.
        - Continue with a back-and-forth conversation, exploring key arguments and points.
        - Include examples and anecdotes as part of the dialogue.
        - Ensure the conversation flows naturally and engagingly.

        📈 {prompts['takeaways']}:
        - List 5-7 practical insights discussed during the conversation.
        - Explain their significance and potential impact.

        Text to summarize: {text}

        Ensure the summary is comprehensive enough for someone who hasn't seen the original content."""

    else:
        system_prompt = f"""You are an expert content analyst and summarizer. Create a comprehensive 
        video-style summary in {target_language}. Ensure all content is fully translated and culturally adapted 
        to the target language."""

        user_prompt = f"""Please provide a detailed video-style summary of the following content in {target_language}. 
        Structure your response as follows:

        🎯 {prompts['title']}: Create a descriptive title

        📝 {prompts['overview']} (2-3 sentences):
        - Provide a brief context and main purpose

        🔑 {prompts['key_points']}:
        - Extract and explain the main arguments
        - Include specific examples
        - Highlight unique perspectives

        💡 {prompts['takeaways']}:
        - List 3-5 practical insights
        - Explain their significance

        🔄 {prompts['context']}:
        - Broader context discussion
        - Future implications

        Text to summarize: {text}

        Ensure the summary is comprehensive enough for someone who hasn't seen the original content."""

    return system_prompt, user_prompt


# Function to create prompts
def create_prompts(language_code, section_number, text_chunk, mode):
    language_instructions = {
        'en': 'Create a detailed summary of the following section in English. Maintain all important information, arguments, and connections.',
        'hi': 'कृपया निम्नलिखित खंड का हिंदी में संक्षिप्त वर्णन करें। सभी महत्वपूर्ण जानकारी, तर्क और कनेक्शन बनाए रखें।',
        'de': 'Erstellen Sie eine detaillierte Zusammenfassung des folgenden Abschnitts auf Deutsch. Behalten Sie alle wichtigen Informationen, Argumente und Verbindungen bei.',
        'it': 'Crea un riassunto dettagliato della seguente sezione in italiano. Mantieni tutte le informazioni importanti, gli argomenti e le connessioni.',
        'es': 'Cree un resumen detallado de la siguiente sección en español. Mantenga toda la información importante, los argumentos y las conexiones.',
        'fr': 'Créez un résumé détaillé de la section suivante en français. Conservez toutes les informations importantes, arguments et connexions.',
        'nl': 'Maak een gedetailleerde samenvatting van het volgende gedeelte in het Nederlands. Behoud alle belangrijke informatie, argumenten en verbindingen.',
        'pl': 'Utwórz szczegółowe podsumowanie następującej sekcji po polsku. Zachowaj wszystkie ważne informacje, argumenty i połączenia.',
        'ja': '以下のセクションの詳細な要約を日本語で作成してください。すべての重要な情報、議論、および接続を維持します。',
        'zh': '用中文创建以下部分的详细摘要。保留所有重要信息、论点和连接。',
        'ru': 'Создайте подробное резюме следующего раздела на русском языке. Сохраните всю важную информацию, аргументы и связи.',
        'ko': '다음 섹션에 대한 자세한 요약을 한국어로 작성하세요. 모든 중요한 정보, 논쟁 및 연결을 유지합니다.',
        'pt': 'Crie um resumo detalhado da seção a seguir em português. Mantenha todas as informações importantes, argumentos e conexões.',
        'ar': 'قم بإنشاء ملخص مفصل للقسم التالي باللغة العربية. حافظ على جميع المعلومات المهمة والحجج والاتصالات.',
        'tr': 'Aşağıdaki bölümün Türkçe ayrıntılı bir özetini oluşturun. Tüm önemli bilgileri, argümanları ve bağlantıları koruyun.',
        'bn': 'নিম্নলিখিত অংশের বাংলায় একটি বিশদ সংক্ষিপ্তসার তৈরি করুন। সমস্ত গুরুত্বপূর্ণ তথ্য, যুক্তি এবং সংযোগগুলি বজায় রাখুন।',
        'mr': 'खालील विभागाचा मराठीत सविस्तर आढावा घ्या. सर्व महत्त्वाची माहिती, तर्क आणि कनेक्शन कायम ठेवा.',
        'ta': 'கீழ்க்கண்ட பகுதியின் தமிழில் விரிவான சுருக்கத்தை உருவாக்குங்கள். அனைத்து முக்கியமான தகவல்களையும் வாதங்களையும் இணைப்புகளையும் பராமரிக்கவும்.',
        'te': 'క్రింది విభాగం యొక్క తెలుగు లో వివరమైన సారాంశాన్ని సృష్టించండి. అన్ని ముఖ్యమైన సమాచారాన్ని, వాదనలను మరియు సంబంధాలను నిర్వహించండి.',
        'kn': 'ಕೆಳಗಿನ ವಿಭಾಗದ ಕನ್ನಡದಲ್ಲಿ ವಿವರವಾದ ಸಂಕ್ಷಿಪ್ತ ವಿವರವನ್ನು ರಚಿಸಿ. ಎಲ್ಲಾ ಮುಖ್ಯ ಮಾಹಿತಿಯನ್ನು, ವಾದಗಳನ್ನು ಮತ್ತು ಸಂಪರ್ಕಗಳನ್ನು ಉಳಿಸಿ.',
        'ml': 'താഴെ പറയുന്ന വിഭാഗത്തിന്റെ മലയാളത്തിൽ വിശദമായ സാരാംശം സൃഷ്ടിക്കുക. എല്ലാ പ്രധാന വിവരങ്ങളും വാദങ്ങളും ബന്ധങ്ങളും നിലനിർത്തുക.',
        'bh': 'निम्नलिखित खंड का भोजपुरी में एक विस्तृत सारांश बनाईं। सभे महत्वपूर्ण जानकारी, तर्क अउर संबंधन के बनावे के काम करीं।'
    }

    instruction = language_instructions.get(language_code, language_instructions['en'])

    if mode == 'podcast':
        instruction += ' Present this summary in a narrative, engaging style suitable for a podcast. Ensure a natural flow with storytelling elements.'
    elif mode == 'video':
        instruction += ' Present this summary in a structured, concise format suitable for a video. Use bullet points and headings to clearly organize the key points.'
    else:
        raise ValueError("Invalid mode. Choose either 'podcast' or 'video'.")

    system_prompt = f"""You are an expert content summarizer. {instruction}
    Ensure the summary is fully in {language_code}, without any words from other languages."""

    user_prompt = f"""{instruction}
    Pay special attention to:
    - Main topics and arguments
    - Important details and examples
    - Connections with other mentioned topics
    - Key statements and conclusions

    Text: {text_chunk}"""

    return system_prompt, user_prompt

# Function to summarize with Langchain and OpenAI
def summarize_with_langchain_and_openai(transcript, mode, language_code='en', model_name='llama-3.1-8b-instant'):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=7000,
        chunk_overlap=1000,
        length_function=len
    )
    texts = text_splitter.split_text(transcript)

    intermediate_summaries = []
    
    for i, text_chunk in enumerate(texts):
        system_prompt, user_prompt = create_summary_prompt(text_chunk, language_code, mode)
        
        try:
            response = groq_client.chat.completions.create(
                model=model_name,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.7,  # Keeping temperature low to maintain consistency
                max_tokens=8000
            )
            
            summary = response.choices[0].message.content
            intermediate_summaries.append(summary)
            
        except Exception as e:
            st.error(f"Error with Groq API during intermediate summarization: {str(e)}")
            return None

    combined_summary = "\n\n=== Next Section ===\n\n".join(intermediate_summaries)
    
    final_instruction = 'Maintain a narrative and engaging style, making sure to connect the points naturally and conversationally. Use transitions and storytelling elements to keep it engaging.' if mode == 'podcast' else 'Keep the summary concise and well-structured, focusing on key points and details. Use bullet points and headings to organize the content clearly.'

    final_system_prompt = f"""You are an expert in creating comprehensive summaries. 
    Create a coherent, well-structured complete summary in {language_code} from the 
    provided intermediate summaries. Connect the information logically and establish 
    important relationships. Ensure the summary is fully in {language_code}, without any words from other languages.
    {final_instruction}"""
    
    final_user_prompt = f"""Create a final, comprehensive summary from the following 
    intermediate summaries. The summary should be fully in {language_code}, without any words from other languages.
    - Include all important topics and arguments
    - Establish logical connections between topics
    - Have a clear structure
    - Highlight key statements and most important insights
    {final_instruction}

    Intermediate summaries:
    {combined_summary}"""
    
    try:
        final_response = groq_client.chat.completions.create(
            model=model_name,
            messages=[
                {"role": "system", "content": final_system_prompt},
                {"role": "user", "content": final_user_prompt}
            ],
            temperature=0.7,  # Keeping temperature low to maintain consistency
            max_tokens=8000
        )
        
        final_summary = final_response.choices[0].message.content
        return final_summary
    except Exception as e:
        st.error(f"Error with Groq API during final summarization: {str(e)}")
        return None


class PDF(FPDF):
    def header(self):
        self.set_font('FreeSerif', 'B', 12)
        self.cell(0, 10, 'Summary Report', 0, 1, 'C')
        self.ln(10)

    def chapter_title(self, title):
        self.set_font('FreeSerif', 'B', 12)
        self.cell(0, 10, title, 0, 1, 'L')
        self.ln(5)

    def chapter_body(self, body):
        self.set_font('FreeSerif', '', 12)
        for line in body.split('\n'):
            if line.strip().endswith(':'):
                self.set_font('FreeSerif', 'B', 12)
                self.multi_cell(0, 10, line.strip())
                self.set_font('FreeSerif', '', 12)
            elif line.startswith('* '):
                self.set_font('FreeSerif', '', 12)
                self.multi_cell(0, 10, u'\u2022 ' + line[2:])
            else:
                self.multi_cell(0, 10, line)
            self.ln(5)

# Function to generate PDF
def generate_pdf(summary, title="Summary"):
    pdf = PDF()

    # Load Unicode fonts from the specified path
    font_path = 'C:\\coding\\youtube summarizer\\youtube_summarizer\\textFormat\\freeserif\\'
    pdf.add_font('FreeSerif', '', font_path + 'FreeSerif.ttf', uni=True)
    pdf.add_font('FreeSerif', 'B', font_path + 'FreeSerifBold.ttf', uni=True)
    pdf.add_font('FreeSerif', 'I', font_path + 'FreeSerifItalic.ttf', uni=True)
    pdf.add_font('FreeSerif', 'BI', font_path + 'FreeSerifBoldItalic.ttf', uni=True)

    pdf.add_page()
    pdf.chapter_title(title)
    pdf.chapter_body(summary)

    # Output PDF to a string
    pdf_output = pdf.output(dest='S').encode('latin1')
    return pdf_output

# Function to generate DOCX
def generate_doc(summary, title="Summary"):
    doc = Document()
    doc.add_heading(title, 0)

    for line in summary.split('\n'):
        if line.startswith('**') and line.endswith('**'):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line.strip('**'))
            run.bold = True
        elif line.strip().endswith(':'):
            doc.add_heading(line.strip(), level=2)
        elif line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)

    with BytesIO() as doc_output:
        doc.save(doc_output)
        doc_output.seek(0)
        return doc_output.read()

# Function to analyze sentiment
def analyze_sentiment(text):
    analyzer = SentimentIntensityAnalyzer()
    sentiment = analyzer.polarity_scores(text)
    return sentiment

def map_key_points_to_intervals(key_points, total_duration):
    """
    Maps key points to evenly distributed time intervals across the video duration.

    Args:
        key_points (list): List of key points (strings).
        total_duration (int): Total duration of the video in seconds.

    Returns:
        list: List of tuples representing (start_time, end_time) for each key point.
    """
    num_points = len(key_points)
    interval_duration = total_duration / num_points  # Divide video into equal segments
    
    intervals = []
    for i in range(num_points):
        start_time = i * interval_duration
        end_time = min((i + 1) * interval_duration, total_duration)
        intervals.append((start_time, end_time))
    
    return intervals

def find_subtitle_index(current_time, intervals):
    """
    Finds the index of the subtitle that corresponds to the given time.

    Args:
        current_time (float): Current time in seconds.
        intervals (list): List of (start_time, end_time) intervals.

    Returns:
        int or None: Index of the subtitle if found, otherwise None.
    """
    for idx, (start_time, end_time) in enumerate(intervals):
        if start_time <= current_time < end_time:
            return idx
    return None

def extract_key_points(summary):
    # Assuming key points are indicated by bullet points or specific markers in the summary
    key_points = [line for line in summary.split('\n') if line.startswith('🔑') or line.startswith('* ')]
    return key_points

def create_highlight_reels(video_path, key_points, subtitles, reel_duration=60):

    # Load video with OpenCV
    cap = cv2.VideoCapture(video_path)
    
    # Get video properties
    fps = int(cap.get(cv2.CAP_PROP_FPS))
    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    total_duration = total_frames / fps  # in seconds
    width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
    height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
    
    # Convert key points to time intervals
    intervals = map_key_points_to_intervals(key_points, total_duration)
    
    # Calculate number of reels
    num_reels = math.ceil(total_duration / reel_duration)
    reel_paths = []

    for i in range(num_reels):
        start_time = i * reel_duration
        end_time = min((i + 1) * reel_duration, total_duration)
        
        output_video_path = f'reel_{i + 1}.mp4'
        fourcc = cv2.VideoWriter_fourcc(*'mp4v')
        out = cv2.VideoWriter(output_video_path, fourcc, fps, (width, height))
        
        cap.set(cv2.CAP_PROP_POS_MSEC, start_time * 1000)
        while cap.isOpened() and cap.get(cv2.CAP_PROP_POS_MSEC) < end_time * 1000:
            ret, frame = cap.read()
            if not ret:
                break
            
            current_time = cap.get(cv2.CAP_PROP_POS_MSEC) / 1000
            subtitle_index = find_subtitle_index(current_time, intervals)

            if subtitle_index is not None:
                current_subtitle = subtitles[subtitle_index]
                cv2.putText(frame, current_subtitle, (10, height - 30), 
                            cv2.FONT_HERSHEY_SIMPLEX, 1, (255, 255, 255), 2, cv2.LINE_AA)
            
            out.write(frame)

        out.release()
        reel_paths.append(output_video_path)
    
    cap.release()
    return reel_paths

def get_time_interval(point):
    # Placeholder for actual mapping of points to time intervals
    time_intervals = {
        "First Key Point": (10, 20),  # start time in seconds, end time in seconds
        "Second Key Point": (30, 45),
        "Third Key Point": (50, 65)
        # Add more mappings as needed
    }
    return time_intervals.get(point.strip(), (0, 5))  # default to the first 5 seconds if not found

def download_youtube_video(url, cookies_path=None):
    try:
        ydl_opts = {
            'format': 'bestaudio+bvideo',  # Download best audio and video streams
            'outtmpl': 'videos/%(title)s.%(ext)s',  # Save video with title
            'quiet': False,  # Show progress during download
        }

        if cookies_path:
            ydl_opts['cookiefile'] = cookies_path

        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info_dict = ydl.extract_info(url, download=True)
            video_path = ydl.prepare_filename(info_dict)
            return video_path
    except Exception as e:
        print(f"Error downloading the video: {str(e)}")
        return None


def main():
    st.title('📺 Advanced YouTube Video Summarizer')
    st.markdown("""
    This tool creates comprehensive summaries of YouTube videos using advanced AI technology.
    It works with both videos that have transcripts and those that don't!
    """)

    # Initialize session state variables
    if 'link' not in st.session_state:
        st.session_state.link = ""
    if 'language' not in st.session_state:
        st.session_state.language = ""
    if 'mode' not in st.session_state:
        st.session_state.mode = ""
    if 'summary' not in st.session_state:
        st.session_state.summary = None
    if 'sentiment' not in st.session_state:
        st.session_state.sentiment = None
    if 'reel' not in st.session_state:
        st.session_state.reel = None

    col1, col2, col3 = st.columns([3, 1, 1])
    
    with col1:
        link = st.text_input('🔗 Enter YouTube video URL:', key='link_input')
    
    with col2:
        languages = get_available_languages()
        target_language = st.selectbox(
            '🌍 Select Summary Language:',
            options=list(languages.keys()),
            index=list(languages.keys()).index(st.session_state.language) if st.session_state.language else 0,
            key='language_input'
        )
        target_language_code = languages[target_language]

    with col3:
        mode = st.selectbox(
            '🎙️ Select Mode:',
            options=['video', 'podcast'],
            index=['video', 'podcast'].index(st.session_state.mode) if st.session_state.mode else 0,
            key='mode_input'
        )
        mode = mode.lower()

    # Check for changes in link, language, or mode
    if (link != st.session_state.link or 
        target_language != st.session_state.language or 
        mode != st.session_state.mode):
        st.session_state.link = link
        st.session_state.language = target_language
        st.session_state.mode = mode
        st.session_state.summary = None
        st.session_state.sentiment = None
        st.session_state.reel = None

    if st.button('Generate Summary'):
        if link:
            try:
                with st.spinner('Processing...'):
                    progress = st.progress(0)
                    status_text = st.empty()

                    status_text.text('📥 Fetching video transcript...')
                    progress.progress(25)

                    transcript, _ = get_transcript(link)

                    status_text.text(f'🤖 Generating {target_language} summary...')
                    progress.progress(75)

                    summary = summarize_with_langchain_and_openai(
                        transcript, 
                        mode,
                        target_language_code,
                        model_name='llama-3.1-8b-instant'
                    )
                    
                    sentiment = analyze_sentiment(transcript)

                    status_text.text('✨ Summary Ready!')
                    # Save summary and sentiment in session state
                    st.session_state.summary = summary
                    st.session_state.sentiment = sentiment

                    progress.progress(100)

            except Exception as e:
                st.error(f"An error occurred: {str(e)}")
        else:
            st.warning('Please enter a valid YouTube link.')

    # Display summary, sentiment, and download buttons if summary exists
    if st.session_state.summary:
        st.markdown(st.session_state.summary)

        sentiment = st.session_state.sentiment
        if sentiment:
            st.markdown("### Sentiment Analysis")
            st.json(sentiment)
            neg = sentiment['neg'] * 100
            neu = sentiment['neu'] * 100
            pos = sentiment['pos'] * 100
            compound = sentiment['compound']
            st.markdown(f"""
            **Negative Sentiment**: {neg:.1f}% - Parts of the text have a negative tone, such as criticism, sadness, or anger.\n
            **Neutral Sentiment**: {neu:.1f}% - The majority of the text is neutral, without strong positive or negative emotions.\n
            **Positive Sentiment**: {pos:.1f}% - Parts of the text have a positive tone, expressing happiness, praise, or optimism.\n
            **Overall Sentiment**: {compound:.4f} - The overall tone of the text is { 'strongly negative' if compound < -0.5 else 'negative' if compound < 0 else 'neutral' if compound == 0 else 'positive' if compound > 0 else 'strongly positive' }.
            """)
            st.markdown(f"""
            **Overall Summary**:
            - **Mostly Neutral**: The majority of the content (about {neu:.1f}%) is neutral, indicating factual or indifferent statements.
            - **Negative Sentiment**: Approximately {neg:.1f}% of the content has a negative tone, suggesting some criticism, sadness, or anger.
            - **Positive Sentiment**: Around {pos:.1f}% of the content is positive, showing instances of happiness, praise, or optimism.
            - **Strong Overall Negative Tone**: The overall sentiment score of {compound:.4f} indicates that the general tone of the video is quite negative.
            """)

        key_points = extract_key_points(st.session_state.summary)
        subtitles = [point.strip('🔑').strip('* ') for point in key_points]

        if st.button('Create Reel(s)'):
            with st.spinner('Creating reel(s)...'):
                video_path = download_youtube_video(link)
                if video_path:
                    subtitles = [point.strip('🔑').strip('* ') for point in key_points]
                    reel_paths = create_highlight_reels(video_path, key_points, subtitles, reel_duration=60)
                    if reel_paths:
                        st.success(f'{len(reel_paths)} reel(s) created successfully!')
                        for i, reel_path in enumerate(reel_paths):
                            st.markdown(f"### Reel {i + 1}")
                            st.download_button(
                                label=f"Download Reel {i + 1} (MP4)",
                                data=open(reel_path, 'rb').read(),
                                file_name=f"reel_{i + 1}.mp4",
                                mime="video/mp4")
                    else:
                        st.error('Failed to create reels.')
                else:
                    st.error('Video download failed. Reel creation aborted.')

        if st.session_state.reel:
            st.markdown('### Download Reel')
            st.download_button(
                label="Download Reel (MP4)",
                data=open(st.session_state.reel, 'rb').read(),
                file_name="reel.mp4",
                mime="video/mp4"
            )

        pdf_data = generate_pdf(st.session_state.summary)
        if pdf_data:
            st.download_button(
                label="Download as PDF",
                data=pdf_data,
                file_name="summary.pdf",
                mime="application/pdf"
            )
        
        doc_data = generate_doc(st.session_state.summary)
        st.download_button(
            label="Download as DOCX",
                data=doc_data,
                file_name="summary.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    main()
